import os
import tempfile
import atexit
import json
from decimal import Decimal
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from PyQt6.QtCore import QObject, pyqtSignal
from PyQt6.QtWidgets import QApplication
from docx.text.paragraph import Paragraph

try:
    import pypandoc
except ImportError:
    pypandoc = None

try:
    from PyQt6.QtWebEngineWidgets import QWebEngineView
except ImportError:
    QWebEngineView = None

from core.constants import TEMPLATE_PATH_MERGED, MATERIAL_CONSUMPTION_MAP
from core.utilities import num_to_words_indian, TEMP_FILES, cleanup_temp_files, OperationCanceledError

atexit.register(cleanup_temp_files)

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '4')
        border_el.set(qn('w:color'), 'auto')
        tblBorders.append(border_el)
    tblPr.append(tblBorders)

def set_table_cell_margins(table, **kwargs):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblCellMar = OxmlElement('w:tblCellMar')
    for m in kwargs:
        if kwargs[m] is not None:
            mar = OxmlElement(f'w:{m}')
            mar.set(qn('w:w'), str(kwargs[m]))
            mar.set(qn('w:type'), 'dxa')
            tblCellMar.append(mar)
    tblPr.append(tblCellMar)

def _generate_abstract_table(document, data):
    items = data.get('items', [])
    if not items:
        return

    # Set page orientation to landscape for this table
    new_section = document.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE

    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('ABSTRACT')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.bold = True
    font.underline = True
    document.add_paragraph() 
    
    headers = ["Item No", "Quantity", "Unit", "Description of Item", "Rate", "Words", "Amount Since Previous", "Amount upto Date"]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    set_table_cell_margins(table, top=40, bottom=0, left=60, right=60)
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.text = header
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs: run.font.bold = True
    
    for item in items:
        row_cells = table.add_row().cells
        rate_val_str = item.get("unit_rate", "0")
        row_cells[0].text, row_cells[1].text, row_cells[2].text = item.get("sr_no", ""), item.get("quantity", ""), item.get("unit", "")
        row_cells[3].text, row_cells[4].text = item.get("description", ""), item.get("unit_rate", "")
        row_cells[5].text = num_to_words_indian(rate_val_str)
        row_cells[6].text, row_cells[7].text = item.get("total", ""), item.get("total", "")
    
    try:
        total_amount_val = float(data.get('total_amount', 'â‚¹0').replace('â‚¹', '').replace(',', ''))
    except (ValueError, TypeError):
        total_amount_val = 0.0
    
    insurance_val = total_amount_val * 0.005
    total_bill_amt_val = total_amount_val + insurance_val

    total_data = [("TOTAL : Rs", total_amount_val), ("Add INSURANCE 0.5 %", insurance_val), ("TOTAL BILL AMT (Rs.)", total_bill_amt_val)]
    for label, value in total_data:
        row_cells = table.add_row().cells
        row_cells[0].merge(row_cells[2])
        p_label = row_cells[3].paragraphs[0]; p_label.add_run(label).bold = True
        row_cells[4].merge(row_cells[5])
        value_str = f"â‚¹{value:,.2f}"
        p_val1 = row_cells[6].paragraphs[0]; p_val1.add_run(value_str).bold = True
        p_val2 = row_cells[7].paragraphs[0]; p_val2.add_run(value_str).bold = True

def _make_table_borderless(table):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_el = OxmlElement(f'w:{border_name}')
                border_el.set(qn('w:val'), 'nil')
                tc_borders.append(border_el)
            tc_pr.append(tc_borders)

def _generate_excess_saving_statement(document, data):
    new_section = document.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE
    
    p_work = document.add_paragraph()
    p_work.add_run("Name of Work\t:\t").bold = True
    p_work.add_run(data.get('name_work', ''))
    
    p_agency = document.add_paragraph()
    p_agency.add_run("Name of Agency\t:\t").bold = True
    p_agency.add_run(data.get('contractor', ''))
    document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('EXCESS SAVING STATEMENT')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.bold = True
    font.underline = True
    document.add_paragraph()
    
    items = data.get('items', [])
    if not items: return

    headers = ["Item No.", "Tender\nQuantity", "Executed\nQuantity", "Unit", "Description of Item", "Excess", "Saving", "Remarks"]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.text = header_text
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs: run.font.bold = True

    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.get("sr_no", "")
        row_cells[1].text = item.get("quantity", "")
        row_cells[2].text = item.get("executed_quantity", "")
        row_cells[3].text = item.get("unit", "")
        row_cells[4].text = item.get("description", "")
        row_cells[5].text = str(item.get("excess", "-"))
        row_cells[6].text = str(item.get("saving", "-"))
        row_cells[7].text = item.get("remarks_excess_saving", "As Per Site Condition")

    document.add_paragraph()
    
    sign_table = document.add_table(rows=1, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.columns[0].width = Inches(3.5)
    sign_table.columns[1].width = Inches(3.5)
    sign_cells = sign_table.rows[0].cells
    
    p_deputy = sign_cells[0].paragraphs[0]
    p_deputy.add_run(data.get('deputy_engineer', 'DEPUTY ENGINEER')).bold = True
    p_deputy.add_run('\nSLUMP IMP. (WEST) SUB DIV NO.')
    p_deputy.add_run('\nM.S.I.BOARD, MHADA, MUMBAI-400051')
    p_deputy.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    p_exec = sign_cells[1].paragraphs[0]
    p_exec.add_run(data.get('executive_engineer', 'EXECUTIVE ENGINEER')).bold = True
    p_exec.add_run('\nSLUMP IMP. (WEST)')
    p_exec.add_run('\nM.S.I.BOARD, MHADA, MUMBAI-400051')
    p_exec.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    _make_table_borderless(sign_table)

def _generate_material_consumption_table(document, data):
    p_work = document.add_paragraph()
    p_work.add_run("Name of Work\t:\t").bold = True
    p_work.add_run(data.get('name_work', ''))
    document.add_paragraph() 
    
    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('MATERIAL CONSUMPTION STATEMENT')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.bold = True
    font.underline = True
    document.add_paragraph() 

    items = data.get('items', [])
    
    consumption_data = []
    material_keys = ["sand", "rubble", "brick", "metal", "cement"]
    consumption_totals = {key: 0.0 for key in material_keys}

    for item in items:
        item_desc_lower = item.get('description', '').lower()
        for keyword, details in MATERIAL_CONSUMPTION_MAP.items():
            if keyword in item_desc_lower:
                try: quantity = float(item.get('quantity', 0))
                except ValueError: quantity = 0.0
                
                row_data = { "item_no": item.get("sr_no"), "short_desc": details["short_desc"], "qty": quantity,
                             "unit": item.get("unit"), "ratios": details["ratios"], "totals": {} }
                for mat_key in material_keys:
                    total_qty = quantity * details["ratios"].get(mat_key, 0.0)
                    row_data["totals"][mat_key] = total_qty
                    consumption_totals[mat_key] += total_qty
                
                consumption_data.append(row_data)
                break
    
    if not consumption_data: return

    table = document.add_table(rows=2, cols=14)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    
    hdr1, hdr2 = table.rows
    hdr1.cells[0].merge(hdr2.cells[0]).text = "Item No"
    hdr1.cells[1].merge(hdr2.cells[1]).text = "Short Description"
    hdr1.cells[2].merge(hdr2.cells[2]).text = "Qty"
    hdr1.cells[3].merge(hdr2.cells[3]).text = "Unit"
    
    mat_details = [("Sand", "M3"), ("Rubble", "M3"), ("Brick", "Nos."), ("Metal", "M3"), ("Cement", "Bags")]
    for i, (name, unit) in enumerate(mat_details):
        hdr1.cells[4 + i*2].merge(hdr1.cells[5 + i*2]).text = name
        hdr2.cells[4 + i*2].text = "Ratio"
        hdr2.cells[5 + i*2].text = f"Total Qty ({unit})"

    for row_data in consumption_data:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data["item_no"]
        row_cells[1].text = row_data["short_desc"]
        row_cells[2].text = f'{row_data["qty"]:.2f}'
        row_cells[3].text = row_data["unit"]
        for i, key in enumerate(material_keys):
            row_cells[4 + i*2].text = f'{row_data["ratios"].get(key, 0.0):.3f}'
            row_cells[5 + i*2].text = f'{row_data["totals"].get(key, 0.0):.2f}'
    
    total_cells = table.add_row().cells
    total_cells[1].text = "Total :"
    total_cells[1].paragraphs[0].runs[0].bold = True
    for i, key in enumerate(material_keys):
        p = total_cells[5 + i*2].paragraphs[0]
        p.add_run(f'{consumption_totals[key]:.2f}').bold = True

def _generate_cement_consumption_table(document, data):
    p_work = document.add_paragraph()
    p_work.add_run("Name of Work\t:\t").bold = True
    p_work.add_run(data.get('name_work', ''))
    
    p_agency = document.add_paragraph()
    p_agency.add_run("Name of Agency\t:\t").bold = True
    p_agency.add_run(data.get('contractor', ''))
    document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('CEMENT CONSUMPTION STATEMENT')
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.bold = True
    font.underline = True
    document.add_paragraph()
    
    items = data.get('items', [])
    cement_items = []
    cement_total = 0.0

    for item in items:
        item_desc_lower = item.get('description', '').lower()
        for keyword, details in MATERIAL_CONSUMPTION_MAP.items():
            if keyword in item_desc_lower and details["ratios"].get("cement", 0.0) > 0:
                try:
                    executed_qty = float(item.get('executed_quantity', '0'))
                    cement_rate = details["ratios"].get("cement", 0.0)
                    theoretical_consumption = executed_qty * cement_rate
                    cement_total += theoretical_consumption
                    cement_items.append({
                        "sr_no": item.get("sr_no", ""),
                        "tender_description": details["short_desc"],
                        "executed_qty": executed_qty,
                        "cement_rate": cement_rate,
                        "unit": item.get("unit", ""),
                        "theoretical_consumption": theoretical_consumption
                    })
                except (ValueError, TypeError):
                    continue
    
    if not cement_items: return

    headers = ["Sr. No", "Tender Description", "Executed\nQuantity", "Rate of\ncement\nConsumption", "Unit", "Theoretical\nConsumption\nin Bag"]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.text = header_text
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs: run.font.bold = True
        
    for item in cement_items:
        row_cells = table.add_row().cells
        row_cells[0].text = item["sr_no"]
        row_cells[1].text = item["tender_description"]
        row_cells[2].text = f'{item["executed_qty"]:.2f}'
        row_cells[3].text = f'{item["cement_rate"]:.3f}'
        row_cells[4].text = item["unit"]
        row_cells[5].text = f'{item["theoretical_consumption"]:.2f}'

    total_row = table.add_row().cells
    total_row[0].merge(total_row[4])
    total_row[0].text = "Total ="
    total_row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    total_row[5].text = f'{cement_total:.2f}'
    
    say_row = table.add_row().cells
    say_row[0].merge(say_row[4])
    say_row[0].text = "Say ="
    say_row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    say_row[5].text = f'{round(cement_total):.0f}'

    document.add_paragraph()
    document.add_paragraph()
    
    sign_table = document.add_table(rows=1, cols=3)
    sign_cells = sign_table.rows[0].cells
    
    p1 = sign_cells[0].paragraphs[0]
    p1.add_run(data.get('executive_engineer', '[Executive Engineer Name]')).bold = True
    p1.add_run("\nSECT ENGINEER/ D.B.")
    p1.add_run("\nEXECUTIVE ENGINEER WEST")
    p1.add_run("\nM.S.I.BOARD, MHADA, MUMBAI-51")
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p2 = sign_cells[1].paragraphs[0]
    p2.add_run(data.get('signatory_jr_engineer', '[Jr. Engineer Name]')).bold = True
    p2.add_run("\nJR./ SECT./ ASST. ENGINEER")
    p2.add_run("\nSLUMP IMP. (WEST) SUB DIV NO")
    p2.add_run("\nM.S.I.BOARD, MHADA, MUMBAI-51")
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p3 = sign_cells[2].paragraphs[0]
    p3.add_run(data.get('deputy_engineer', '[Deputy Engineer Name]')).bold = True
    p3.add_run("\nDEPUTY ENGINEER")
    p3.add_run("\nSLUMP IMP. (WEST) SUB DIV NO")
    p3.add_run("\nM.S.I.BOARD, MHADA, MUMBAI-51")
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    _make_table_borderless(sign_table)

def generate_merged_form_report(data, output_path, template_path):
    if not os.path.exists(template_path):
        return False, f"Template file not found: {template_path}"

    try:
        document = Document(template_path)

        for p in document.paragraphs:
            for key, val in data.items():
                if isinstance(val, (str, int, float)):
                    inline = p.runs
                    full_text = "".join(run.text for run in inline)
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in full_text:
                        new_text = full_text.replace(placeholder, str(val))
                        p.clear()
                        p.add_run(new_text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, val in data.items():
                            if isinstance(val, (str, int, float)):
                                inline = p.runs
                                full_text = "".join(run.text for run in inline)
                                placeholder = f"{{{{{key}}}}}"
                                if placeholder in full_text:
                                    new_text = full_text.replace(placeholder, str(val))
                                    p.clear()
                                    p.add_run(new_text)

        all_paragraphs = list(document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        for p in all_paragraphs:
            if '{{abstract_table}}' in "".join(r.text for r in p.runs):
                p.clear()
                _generate_abstract_table(document, data)
                break
        
        for p in all_paragraphs:
            if '{{excess_saving_statement_table}}' in "".join(r.text for r in p.runs):
                p.clear()
                _generate_excess_saving_statement(document, data)
                break
        
        for p in all_paragraphs:
            if '{{material_consumption_statement_table}}' in "".join(r.text for r in p.runs):
                p.clear()
                _generate_material_consumption_table(document, data)
                break
        
        for p in all_paragraphs:
            if '{{cement_consumption_statement_table}}' in "".join(r.text for r in p.runs):
                p.clear()
                _generate_cement_consumption_table(document, data)
                break
        
        document.save(output_path)
        return True, None
    except Exception as e:
        import traceback
        traceback.print_exc()
        return False, f"Failed to generate DOCX: {e}"

def generate_docx_internal(data, output_path):
    try:
        return generate_merged_form_report(data, output_path, TEMPLATE_PATH_MERGED)
    except Exception as e:
        return False, f"Document generation failed: {e}"

def convert_docx_to_pdf(data, output_pdf_path):
    docx_temp_fd, docx_temp_path = tempfile.mkstemp(suffix=".docx")
    os.close(docx_temp_fd)
    TEMP_FILES.append(docx_temp_path)

    success_docx, msg_docx = generate_docx_internal(data, docx_temp_path)
    if not success_docx:
        return False, msg_docx

    try:
        if not pypandoc:
            return False, "pypandoc library is not installed."
        
        pypandoc_extra_args = [
            '--pdf-engine=xelatex',
            '-V', 'mainfont=Arial'
        ]
            
        pypandoc.convert_file(docx_temp_path, 'pdf', outputfile=output_pdf_path, extra_args=pypandoc_extra_args)
        return True, None
    except RuntimeError as e:
        error_msg = (
            "PDF conversion failed. Please ensure Pandoc and a LaTeX engine "
            "(like MiKTeX with xelatex) are installed and in your system's PATH.\n\n"
            f"Details: {e}"
        )
        return False, error_msg
    finally:
        if os.path.exists(docx_temp_path):
            try:
                os.remove(docx_temp_path)
            except OSError:
                pass

def generate_html_preview(data):
    styles = """
    <style>
        body { font-family: Arial, sans-serif; font-size: 10pt; background-color: #f8f8f8; color: #333; }
        .page { background-color: white; padding: 40px; margin: 20px auto; max-width: 800px; box-shadow: 0 0 10px rgba(0,0,0,0.1); }
        h3 { text-align: center; font-weight: bold; text-decoration: underline; }
        h4 { text-align: center; font-weight: bold; }
        table { border-collapse: collapse; width: 100%; font-size: 9pt; margin-top: 15px; }
        th, td { border: 1px solid black; padding: 4px; text-align: left; vertical-align: top; }
        th { font-weight: bold; text-align: center; background-color: #e0e0e0; }
        .no-border, .no-border td { border: none; }
        .header-info { margin-bottom: 15px; }
        .header-info p { margin: 2px 0; }
        .signatory-block { display: inline-block; width: 30%; text-align: center; vertical-align: top; margin-top: 30px; }
        .letter-body { line-height: 1.6; }
        .letter-header { display: flex; justify-content: space-between; }
        .letter-header .right { text-align: left; }
    </style>
    """
    body = ""
    body += "<div class='page'>"
    body += "<div class='letter-body'>"
    body += f"<div class='letter-header'><div class='left'><p><b>Fund Head:</b> {data.get('fund_head', '')}<br><b>Name:</b> {data.get('name', '')}<br><b>Constituency:</b> {data.get('constituency', '')}</p></div><div class='right'><p>Office of the Deputy Engineer<br>M.S.I.B. WEST Division<br>MHADA, Bandra (E),<br>Mumbai-400051.</p></div></div>"
    body += f"<p><b>To,</b><br>{data.get('send_to', '')}<br>M.S.I.B. WEST Division<br>MHADA, Mumbai.</p>"
    body += f"<p><b>Sub: Submission of {data.get('subject', '')}</b></p>"
    body += f"<p><b>Sir,</b><br>I am submitting herewith the {data.get('message', '')} of above work along with site statement & M.B.No. {data.get('mb_no', '')} for making payment to the contractor {data.get('contractor', '')}.</p>"
    body += f"<p><b>Agreement No:</b> {data.get('agreement_no', '')}</p>"
    body += "<p>Yours faithfully,</p>"
    body += f"<p><br><b>({data.get('deputy_engineer', 'Deputy Engineer')})</b><br>M.S.I.B. WEST Division<br>MHADA, Mumbai.</p>"
    body += f"<p>D.A.: M.B.No. {data.get('mb_no', '')}</p>"
    body += "</div></div>"
    body += "<div class='page'>"
    body += "<div class='letter-body'>"
    body += f"<div class='letter-header'><div class='left'><p><b>Fund Head:</b> {data.get('fund_head', '')}<br><b>Name:</b> {data.get('name', '')}<br><b>Constituency:</b> {data.get('constituency', '')}</p></div><div class='right'><p>Office of the Executive Engineer<br>M.S.I.B. WEST Division<br>MHADA, Bandra (E),<br>Mumbai-400051.</p></div></div>"
    body += f"<p><b>To,</b><br>{data.get('send_to', '')}<br>M.S.I. Board, Mumbai.</p>"
    body += f"<p><b>Sub: Submission of {data.get('subject', '')}</b></p>"
    body += f"<p><b>Sir,</b><br>I am submitting herewith the {data.get('message', '')} of above work along with site statement & M.B.No. {data.get('mb_no', '')} for making payment to the contractor {data.get('contractor', '')}.</p>"
    body += f"<p><b>Agreement No:</b> {data.get('agreement_no', '')}</p>"
    body += "<p>Yours faithfully,</p>"
    body += f"<p><br><b>({data.get('deputy_engineer', 'Deputy Engineer')})</b><br>M.S.I.B. WEST Division<br>MHADA, Mumbai.</p>"
    body += f"<p>D.A.: M.B.No. {data.get('mb_no', '')}</p>"
    body += "</div></div>"
    body += "<div class='page'><h3>FORM 47</h3><h4>RUNNING ACCOUNT BILL</h4>"
    body += "<table class='no-border'><tr><td>Division: MSIB West Division</td><td></td></tr><tr><td>Sub-Division: Sub Division No.</td><td></td></tr></table>"
    body += f"<table><tr><td colspan='2'>Name of Contractor: {data.get('contractor', '')}</td><td colspan='2'>Serial No. of this bill: {data.get('message', '')}</td></tr>"
    body += f"<tr><td colspan='2'>Name of Work: {data.get('name_work', '')}</td><td colspan='2'>No. and date of previous bill:</td></tr>"
    body += f"<tr><td colspan='2'>Reference to agreement: {data.get('agreement_no', '')}</td><td colspan='2'>Acceptance No: {data.get('acceptance_no', '')} &nbsp;&nbsp; Date: {data.get('date', '')}</td></tr>"
    body += f"<tr><td colspan='2'>Work Order No: {data.get('work_order_no', '')}</td><td colspan='2'>Date of written order to commence work: {data.get('date', '')}</td></tr>"
    body += f"<tr><td colspan='2'>Date of completion stipulated in contract: {data.get('end_date', '')}</td><td colspan='2'>Date of actual completion of work:</td></tr></table></div>"
    body += "<div class='page'><h3>Annexure – I</h3>"
    body += f"<p><b>Name of Work:</b> {data.get('name_work', '')}<br>"
    body += f"<b>Fund Head:</b> {data.get('fund_head', '')}<br>"
    body += f"<b>Constituency:</b> {data.get('constituency', '')}</p>"
    body += f"<b>Name of Agency:</b> {data.get('contractor', '')}<br>"
    body += f"<b>Agreement No:</b> {data.get('agreement_no', '')}</p>"
    body += "<h4>CERTIFICATE</h4><ol style='list-style-position: inside; padding-left: 0;'>"
    body += "<li>Materials are used in subjected are as per specifications.</li>"
    body += "<li>Construction material has been tested and test reports are found satisfactory.</li>"
    body += "<li>The subjected site is not inspected by Vigilance and Quality Control Cell / A and hence the question of pending remarks does not arise.</li>"
    body += "<li>Nothing is outstanding against the contractor.</li>"
    body += "<li>It is to certify that the contractors have not put any sort of claim against the subjected work.</li></ol>"
    body += f"""
    <div style='width: 100%; margin-top: 40px;'>
        <div class='signatory-block'><b>J.E./S.E./Asst. Engineer</b><br>M.S.I.B. West Div</div>
        <div class='signatory-block'><b>Dy. Engineer</b><br>M.S.I.B. West Div</div>
        <div class='signatory-block'><b>Executive Engineer</b><br>M.S.I.B. West Div</div>
    </div></div>"""
    body += "<div class='page'><h3>Check List to be Attached with Bills of Contractor</h3>"
    body += "<table>"
    body += f"<tr><td>1</td><td>Name of Work</td><td>:</td><td>{data.get('name_work', '')}</td></tr>"
    body += f"<tr><td>2</td><td>Administrative Approval Accorded by the collector</td><td>:</td><td>Amount Rs. {data.get('amt_rupes', '')} <br>Letter No. {data.get('letter_no', '')} <br>Date: {data.get('date', '')}</td></tr>"
    body += f"<tr><td>3</td><td>Technical Sanction accorded by Executive Engineer</td><td>:</td><td>Vide letter No: {data.get('vide_letter_no', '')} Date: {data.get('date', '')}<br>Amount Rs: {data.get('amt_rupes', '')}<br>In Year: {data.get('year', '')}</td></tr>"
    body += f"<tr><td>4</td><td>Estimated cost put to tender</td><td>:</td><td>{data.get('est_cost', '')}</td></tr>"
    body += f"<tr><td>5</td><td>Name of Agency</td><td>:</td><td>{data.get('contractor', '')}</td></tr>"
    body += f"<tr><td>6</td><td>Percentage Quoted</td><td>:</td><td>{data.get('percentage_quoted', '')}</td></tr>"
    body += f"<tr><td>8</td><td>Agreement No</td><td>:</td><td>{data.get('agreement_no', '')}</td></tr>"
    body += f"<tr><td>9</td><td>Date of start of work</td><td>:</td><td>{data.get('start_date', '')}</td></tr>"
    body += f"<tr><td>10</td><td>Stipulated date of completion</td><td>:</td><td>{data.get('end_date', '')}</td></tr></table></div>"
    body += "<div class='page'><h3>ABSTRACT</h3>"
    if data.get('items'):
        body += "<table><tr><th>Item No</th><th>Quantity</th><th>Unit</th><th>Description of Item</th><th>Rate</th><th>Amount upto Date</th></tr>"
        for item in data.get('items', []):
            body += f"<tr><td>{item.get('sr_no', '')}</td><td>{item.get('quantity', '')}</td><td>{item.get('unit', '')}</td><td>{item.get('description', '')}</td><td>{item.get('unit_rate', '')}</td><td>{item.get('total', '')}</td></tr>"
        try:
            total_amount_val = float(data.get('total_amount', 'â‚¹0').replace('â‚¹', '').replace(',', ''))
        except (ValueError, TypeError): total_amount_val = 0.0
        insurance_val = total_amount_val * 0.005
        total_bill_amt_val = total_amount_val + insurance_val
        body += f"<tr><td colspan='4' style='text-align:right;'><b>TOTAL : Rs</b></td><td colspan='2'><b>â‚¹{total_amount_val:,.2f}</b></td></tr>"
        body += f"<tr><td colspan='4' style='text-align:right;'><b>Add INSURANCE 0.5 %</b></td><td colspan='2'><b>â‚¹{insurance_val:,.2f}</b></td></tr>"
        body += f"<tr><td colspan='4' style='text-align:right;'><b>TOTAL BILL AMT (Rs.)</b></td><td colspan='2'><b>â‚¹{total_bill_amt_val:,.2f}</b></td></tr></table>"
    body += f"""
    <div style='width: 100%; margin-top: 40px;'>
        <div class='signatory-block'><b>J.E./S.E./Asst. Engineer</b><br>M.S.I.B. West Div</div>
        <div class='signatory-block'><b>Dy. Engineer</b><br>M.S.I.B. West Div</div>
        <div class='signatory-block'><b>Executive Engineer</b><br>M.S.I.B. West Div</div>
    </div></div>"""
    body += "<div class='page'><h3>MATERIAL CONSUMPTION STATEMENT</h3>"
    items = data.get('items', [])
    consumption_data = []
    material_keys = ["sand", "rubble", "brick", "metal", "cement"]
    consumption_totals = {key: 0.0 for key in material_keys}
    for item in items:
        item_desc_lower = item.get('description', '').lower()
        for keyword, details in MATERIAL_CONSUMPTION_MAP.items():
            if keyword in item_desc_lower:
                try: quantity = float(item.get('quantity', 0))
                except ValueError: quantity = 0.0
                row_data = { "item_no": item.get("sr_no", ""), "short_desc": details["short_desc"], "qty": quantity, "unit": item.get("unit"), "ratios": details["ratios"], "totals": {} }
                for mat_key in material_keys:
                    total_qty = quantity * details["ratios"].get(mat_key, 0.0)
                    row_data["totals"][mat_key] = total_qty
                    consumption_totals[mat_key] += total_qty
                consumption_data.append(row_data)
                break
    if consumption_data:
        body += "<table>"
        body += "<tr><th rowspan='2'>Item No</th><th>Description</th><th rowspan='2'>Qty</th><th rowspan='2'>Unit</th>"
        body += "<th colspan='2'>Sand</th><th colspan='2'>Rubble</th><th colspan='2'>Brick</th><th colspan='2'>Metal</th><th colspan='2'>Cement</th></tr>"
        body += "<tr><th>Ratio</th><th>Total Qty (M3)</th><th>Ratio</th><th>Total Qty (M3)</th><th>Ratio</th><th>Total Qty (Nos.)</th><th>Ratio</th><th>Total Qty (M3)</th><th>Ratio</th><th>Total Qty (Bags)</th></tr>"
        for row in consumption_data:
            body += f"<tr><td>{row['item_no']}</td><td>{row['short_desc']}</td><td>{row['qty']:.2f}</td><td>{row['unit']}</td>"
            for key in material_keys:
                body += f"<td>{row['ratios'].get(key, 0.0):.3f}</td><td>{row['totals'].get(key, 0.0):.2f}</td>"
            body += "</tr>"
        body += "<tr><td colspan='2' style='text-align:right;'><b>Total:</b></td><td></td><td></td>"
        for key in material_keys:
            body += f"<td></td><td style='font-weight:bold;'>{consumption_totals[key]:.2f}</td>"
        body += "</tr></table>"
    body += "</div>"
    body += "<div class='page'><h3>EXCESS SAVING STATEMENT</h3>"
    if data.get('items'):
        body += "<table><tr>"
        headers = ["Item No.", "Tender Qty", "Executed Qty", "Unit", "Description", "Excess", "Saving", "Remarks"]
        for h in headers:
            body += f"<th>{h}</th>"
        body += "</tr>"
        for item in data.get('items', []):
            body += "<tr>"
            body += f"<td>{item.get('sr_no', '')}</td><td>{item.get('quantity', '')}</td><td>{item.get('executed_quantity', '')}</td><td>{item.get('unit', '')}</td><td>{item.get('description', '')}</td><td>{item.get('excess', '-')}</td><td>{item.get('saving', '-')}</td><td>{item.get('remarks_excess_saving', 'As Per Site Condition')}</td>"
            body += "</tr>"
        body += "</table>"
    body += "</div>"
    
    return f"<html><head>{styles}</head><body>{body}</body></html>"

class DocGenWorker(QObject):
    finished = pyqtSignal(bool, str, str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.data = {}
        self.action_type = ""
        self.output_path = ""

    def run_job(self, data, action_type, output_path=""):
        self.data = data
        self.action_type = action_type
        self.output_path = output_path

        try:
            if self.action_type == "fast_preview":
                html_content = generate_html_preview(self.data)
                self.finished.emit(True, "Preview generated.", html_content)

            elif self.action_type == "save_docx":
                success, msg = generate_docx_internal(self.data, self.output_path)
                self.finished.emit(success, msg, self.output_path)
            
            elif self.action_type == "save_pdf":
                success, msg = convert_docx_to_pdf(self.data, self.output_path)
                self.finished.emit(success, msg, self.output_path)

            elif self.action_type == "preview":
                docx_temp_path = tempfile.mkstemp(suffix=".docx", prefix="preview_")[1]
                TEMP_FILES.append(docx_temp_path)
                success_docx, msg_docx = generate_docx_internal(self.data, docx_temp_path)

                if not success_docx:
                    self.finished.emit(False, msg_docx, "")
                    return

                if QWebEngineView and pypandoc:
                    html_temp_path = tempfile.mkstemp(suffix=".html", prefix="preview_")[1]
                    TEMP_FILES.append(html_temp_path)
                    try:
                        pypandoc.convert_file(docx_temp_path, 'html', outputfile=html_temp_path, extra_args=['--mathml'])
                        self.finished.emit(True, "Preview generated successfully.", html_temp_path)
                    except RuntimeError as e:
                        self.finished.emit(False, f"Pandoc HTML conversion failed. Please ensure Pandoc is installed and in your PATH. Error: {e}", "")
                else:
                    self.finished.emit(False, "In-software preview not available. Please ensure pandoc and PyQt6-WebEngine are installed.", "")

            else:
                self.finished.emit(False, "Invalid action type.", "")
        except Exception as e:
            self.finished.emit(False, f"Unexpected error: {str(e)}", "")